"""
Document parser — ekstrak teks dari PDF, DOCX, dan TXT.
"""
import io
from typing import Optional


def parse_file(filename: str, content: bytes) -> Optional[str]:
    """
    Ekstrak teks dari berbagai format file.
    Return teks yang diekstrak, atau None jika gagal.
    """
    ext = filename.lower().split(".")[-1]
    
    try:
        if ext == "txt":
            return content.decode("utf-8", errors="ignore")
        
        elif ext == "pdf":
            try:
                import PyPDF2
                reader = PyPDF2.PdfReader(io.BytesIO(content))
                text = ""
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                # Jika PyPDF2 gagal ekstrak teks, coba pdfplumber
                if not text.strip():
                    try:
                        import pdfplumber
                        with pdfplumber.open(io.BytesIO(content)) as pdf:
                            for page in pdf.pages:
                                text += page.extract_text() or ""
                    except ImportError:
                        pass
                return text.strip() if text.strip() else None
            except ImportError:
                return None
        
        elif ext in ("doc", "docx"):
            try:
                import docx
                doc = docx.Document(io.BytesIO(content))
                text = "\n".join(para.text for para in doc.paragraphs if para.text)
                # Juga ekstrak dari tabel jika ada
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text for cell in row.cells)
                        if row_text.strip():
                            text += "\n" + row_text
                return text.strip() if text.strip() else None
            except ImportError:
                return None
        
        elif ext in ("md", "markdown"):
            return content.decode("utf-8", errors="ignore")
        
        elif ext == "csv":
            text = content.decode("utf-8", errors="ignore")
            # Limit CSV to first 2000 lines
            lines = text.split("\n")[:2000]
            return "\n".join(lines)
        
        else:
            # Coba decode sebagai teks biasa untuk format tidak dikenal
            try:
                return content.decode("utf-8", errors="ignore")
            except:
                return None
    
    except Exception as e:
        print(f"parse_file error ({ext}): {e}")
        return None
